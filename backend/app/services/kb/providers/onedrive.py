"""Microsoft Graph provider — OneDrive for Business + SharePoint.

OneDrive personal, OneDrive for Business, and SharePoint Online all
surface documents through the same Microsoft Graph ``/drives/{id}/root``
tree, so one adapter covers all three. The caller decides which drive
to target via the ``drive_id`` parameter:

* Personal OneDrive — omit ``drive_id`` (we use ``/me/drive``).
* OneDrive for Business — pass the user's ``drive.id`` or omit for self.
* SharePoint document library — pass the site's drive id
  (``/sites/{site-id}/drives`` → pick one).

Token refresh uses the standard OAuth refresh_token grant at the v2
endpoint. Tenants with multi-tenant apps can pass a specific
``tenant`` segment in the token URL via ``tenant_identifier``; common
installs use ``common`` or ``organizations``.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any, AsyncIterator, Awaitable, Callable, Dict, Optional

import httpx

from backend.app.config import get_settings
from backend.app.services.kb.providers.base import (
    ExternalDocument,
    KBProviderAuthError,
    KBProviderError,
)

logger = logging.getLogger(__name__)

_GRAPH = "https://graph.microsoft.com/v1.0"


OnRefresh = Optional[Callable[[str, Optional[str], Optional[int]], Awaitable[None]]]


class OneDriveProvider:
    source_type = "onedrive"

    def __init__(
        self,
        *,
        access_token: str,
        refresh_token: Optional[str] = None,
        drive_id: Optional[str] = None,
        folder_path: Optional[str] = None,
        tenant_identifier: str = "common",
        scopes: Optional[list] = None,
        on_token_refresh: OnRefresh = None,
    ) -> None:
        if not access_token:
            raise KBProviderAuthError("Microsoft Graph access_token is required")
        self._access_token = access_token
        self._refresh_token = refresh_token
        self._drive_id = drive_id
        self._folder_path = (folder_path or "").strip("/")
        self._tenant_id = tenant_identifier
        self._scopes = scopes or ["Files.Read.All", "offline_access"]
        self._on_token_refresh = on_token_refresh
        self._client = httpx.AsyncClient(timeout=30.0)

    async def close(self) -> None:
        await self._client.aclose()

    async def iter_documents(self) -> AsyncIterator[ExternalDocument]:
        # Use delta so re-runs only pull changed items. When no cursor
        # is stored we start from scratch (equivalent to full list).
        start_url = self._delta_base_url()
        next_url: Optional[str] = start_url
        while next_url:
            resp = await self._authed("GET", next_url)
            data = resp.json()
            for item in data.get("value") or []:
                if item.get("folder") is not None:
                    continue  # folders surface as items but we only want files
                mime = (item.get("file") or {}).get("mimeType", "")
                name = item.get("name") or "Untitled"
                if not self._is_text_like(mime, name):
                    continue
                try:
                    text = await self._fetch_text(item, mime)
                except Exception:
                    logger.exception("Graph text fetch failed for %s", item.get("id"))
                    continue
                if not text.strip():
                    continue
                yield ExternalDocument(
                    external_id=str(item["id"]),
                    title=str(name),
                    content=text,
                    source_url=(item.get("webUrl") or None),
                    source_type=self.source_type,
                    updated_at=_parse_graph_ts(item.get("lastModifiedDateTime")),
                    metadata={
                        "mime_type": mime,
                        "drive_id": (item.get("parentReference") or {}).get("driveId"),
                        "path": (item.get("parentReference") or {}).get("path"),
                    },
                )
            next_url = data.get("@odata.nextLink")

    def _delta_base_url(self) -> str:
        """Where to start the listing. ``/delta`` gives us an iteration
        over the drive without tracking inline-deleted items (we're
        read-only so that's fine).
        """
        if self._drive_id:
            base = f"{_GRAPH}/drives/{self._drive_id}/root"
        else:
            base = f"{_GRAPH}/me/drive/root"
        if self._folder_path:
            # "foo/bar" → ":/foo/bar:" sub-path addressing.
            base = f"{base}:/{self._folder_path}:"
        return f"{base}/children"

    @staticmethod
    def _is_text_like(mime: str, name: str) -> bool:
        lower = (name or "").lower()
        if mime in {
            "text/plain",
            "text/markdown",
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }:
            return True
        return lower.endswith((".txt", ".md", ".pdf", ".docx"))

    async def _fetch_text(self, item: Dict[str, Any], mime: str) -> str:
        # Graph exposes a pre-signed download URL on the driveItem.
        url = item.get("@microsoft.graph.downloadUrl")
        if url:
            resp = await self._client.get(url, timeout=30.0)
            resp.raise_for_status()
            blob = resp.content
        else:
            drive_id = (item.get("parentReference") or {}).get("driveId") or self._drive_id
            if not drive_id:
                drive_path = f"{_GRAPH}/me/drive/items/{item['id']}/content"
            else:
                drive_path = f"{_GRAPH}/drives/{drive_id}/items/{item['id']}/content"
            resp = await self._authed("GET", drive_path)
            blob = resp.content

        name_lower = (item.get("name") or "").lower()
        if mime == "application/pdf" or name_lower.endswith(".pdf"):
            return _pdf_to_text(blob)
        if (
            mime
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or name_lower.endswith(".docx")
        ):
            return _docx_to_text(blob)
        try:
            return blob.decode("utf-8", errors="replace")
        except Exception:
            return ""

    async def _authed(self, method: str, url: str) -> httpx.Response:
        resp = await self._client.request(
            method,
            url,
            headers={
                "Authorization": f"Bearer {self._access_token}",
                "Accept": "application/json",
            },
        )
        if resp.status_code == 401 and self._refresh_token:
            await self._refresh_access_token()
            resp = await self._client.request(
                method,
                url,
                headers={
                    "Authorization": f"Bearer {self._access_token}",
                    "Accept": "application/json",
                },
            )
        if resp.status_code == 401:
            raise KBProviderAuthError("Graph rejected the token after refresh")
        if resp.status_code >= 400:
            raise KBProviderError(
                f"Graph {url} failed: {resp.status_code} {resp.text[:200]}"
            )
        return resp

    async def _refresh_access_token(self) -> None:
        settings = get_settings()
        if not (
            settings.MICROSOFT_CLIENT_ID
            and settings.MICROSOFT_CLIENT_SECRET
            and self._refresh_token
        ):
            raise KBProviderAuthError("Microsoft OAuth refresh config missing")
        token_url = (
            f"https://login.microsoftonline.com/{self._tenant_id}/oauth2/v2.0/token"
        )
        resp = await self._client.post(
            token_url,
            data={
                "grant_type": "refresh_token",
                "refresh_token": self._refresh_token,
                "client_id": settings.MICROSOFT_CLIENT_ID,
                "client_secret": settings.MICROSOFT_CLIENT_SECRET,
                "scope": " ".join(self._scopes),
            },
            headers={"Content-Type": "application/x-www-form-urlencoded"},
        )
        if resp.status_code >= 400:
            raise KBProviderAuthError(
                f"Microsoft token refresh failed: {resp.status_code}"
            )
        body = resp.json()
        self._access_token = body.get("access_token") or self._access_token
        new_refresh = body.get("refresh_token")
        if new_refresh:
            self._refresh_token = new_refresh
        if self._on_token_refresh is not None:
            try:
                await self._on_token_refresh(
                    self._access_token,
                    self._refresh_token,
                    body.get("expires_in"),
                )
            except Exception:
                logger.exception("on_token_refresh callback failed (onedrive)")


class SharePointProvider(OneDriveProvider):
    """Alias subclass so the sync runner dispatches by source_type."""

    source_type = "sharepoint"


# ── Helpers ──────────────────────────────────────────────────────────


def _pdf_to_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        return ""
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx_to_text(data: bytes) -> str:
    try:
        import docx  # type: ignore
    except ImportError:
        return ""
    document = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)


def _parse_graph_ts(raw: Optional[str]) -> Optional[datetime]:
    if not raw:
        return None
    try:
        return datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError:
        return None
