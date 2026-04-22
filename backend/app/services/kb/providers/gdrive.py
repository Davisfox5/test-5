"""Google Drive / Google Docs provider.

Enumerates files the tenant's OAuth token can see (scoped to
``https://www.googleapis.com/auth/drive.readonly`` at minimum) and
yields their text. Native Google Docs are exported as ``text/plain``
via the Drive ``files.export`` endpoint; PDFs and DOCX are downloaded
and passed through pypdf / python-docx.

OAuth token refresh is handled inline against Google's token endpoint.
New tokens flow out through the ``on_token_refresh`` callback so the
Integration row stays fresh without a second round-trip.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any, AsyncIterator, Awaitable, Callable, Dict, Optional

import httpx

from backend.app.config import get_settings
from backend.app.services.kb.providers.base import (
    ExternalDocument,
    KBProviderAuthError,
    KBProviderError,
)

logger = logging.getLogger(__name__)

_TOKEN_URL = "https://oauth2.googleapis.com/token"
_DRIVE_LIST = "https://www.googleapis.com/drive/v3/files"
_DRIVE_EXPORT = "https://www.googleapis.com/drive/v3/files/{file_id}/export"
_DRIVE_DOWNLOAD = "https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"

# MIME types we can extract text from. Anything else is skipped.
_NATIVE_DOCS = {"application/vnd.google-apps.document"}
_BLOB_DOCS = {
    "application/pdf",
    "text/plain",
    "text/markdown",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


OnRefresh = Optional[Callable[[str, Optional[str], Optional[int]], Awaitable[None]]]


class GoogleDriveProvider:
    source_type = "gdrive"

    def __init__(
        self,
        *,
        access_token: str,
        refresh_token: Optional[str] = None,
        folder_id: Optional[str] = None,
        drive_id: Optional[str] = None,
        on_token_refresh: OnRefresh = None,
    ) -> None:
        if not access_token:
            raise KBProviderAuthError("Google Drive access_token is required")
        self._access_token = access_token
        self._refresh_token = refresh_token
        self._folder_id = folder_id
        self._drive_id = drive_id
        self._on_token_refresh = on_token_refresh
        self._client = httpx.AsyncClient(timeout=30.0)

    async def close(self) -> None:
        await self._client.aclose()

    async def iter_documents(self) -> AsyncIterator[ExternalDocument]:
        page_token: Optional[str] = None
        while True:
            params: Dict[str, Any] = {
                "pageSize": 100,
                "fields": "nextPageToken, files(id, name, mimeType, webViewLink, modifiedTime, parents)",
                "orderBy": "modifiedTime desc",
            }
            # Shared-drive scope: when a ``drive_id`` is set we have to
            # tell the API to look beyond "My Drive".
            if self._drive_id:
                params.update(
                    {
                        "corpora": "drive",
                        "driveId": self._drive_id,
                        "includeItemsFromAllDrives": "true",
                        "supportsAllDrives": "true",
                    }
                )
            q_parts: list[str] = ["trashed = false"]
            if self._folder_id:
                q_parts.append(f"'{self._folder_id}' in parents")
            params["q"] = " and ".join(q_parts)
            if page_token:
                params["pageToken"] = page_token

            resp = await self._authed("GET", _DRIVE_LIST, params=params)
            data = resp.json()
            for f in data.get("files") or []:
                mime = f.get("mimeType", "")
                if mime not in _NATIVE_DOCS and mime not in _BLOB_DOCS:
                    continue
                try:
                    text = await self._fetch_text(f["id"], mime)
                except Exception:
                    logger.exception("GDrive text fetch failed for %s", f.get("id"))
                    continue
                if not text.strip():
                    continue
                yield ExternalDocument(
                    external_id=str(f["id"]),
                    title=str(f.get("name") or "Untitled"),
                    content=text,
                    source_url=f.get("webViewLink"),
                    source_type=self.source_type,
                    updated_at=_parse_rfc3339(f.get("modifiedTime")),
                    metadata={"mime_type": mime},
                )

            page_token = data.get("nextPageToken")
            if not page_token:
                break

    async def _fetch_text(self, file_id: str, mime_type: str) -> str:
        if mime_type in _NATIVE_DOCS:
            url = _DRIVE_EXPORT.format(file_id=file_id)
            resp = await self._authed(
                "GET", url, params={"mimeType": "text/plain"}
            )
            return resp.text

        url = _DRIVE_DOWNLOAD.format(file_id=file_id)
        resp = await self._authed("GET", url)
        if mime_type == "application/pdf":
            return _pdf_to_text(resp.content)
        if (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return _docx_to_text(resp.content)
        # text/plain or text/markdown
        return resp.text

    async def _authed(
        self, method: str, url: str, *, params: Optional[Dict[str, Any]] = None
    ) -> httpx.Response:
        resp = await self._client.request(
            method,
            url,
            params=params,
            headers={"Authorization": f"Bearer {self._access_token}"},
        )
        if resp.status_code == 401 and self._refresh_token:
            await self._refresh_access_token()
            resp = await self._client.request(
                method,
                url,
                params=params,
                headers={"Authorization": f"Bearer {self._access_token}"},
            )
        if resp.status_code == 401:
            raise KBProviderAuthError("Google Drive rejected the token after refresh")
        if resp.status_code >= 400:
            raise KBProviderError(
                f"GDrive {url} failed: {resp.status_code} {resp.text[:200]}"
            )
        return resp

    async def _refresh_access_token(self) -> None:
        settings = get_settings()
        if not (
            settings.GOOGLE_CLIENT_ID
            and settings.GOOGLE_CLIENT_SECRET
            and self._refresh_token
        ):
            raise KBProviderAuthError("Google OAuth refresh config missing")
        resp = await self._client.post(
            _TOKEN_URL,
            data={
                "grant_type": "refresh_token",
                "refresh_token": self._refresh_token,
                "client_id": settings.GOOGLE_CLIENT_ID,
                "client_secret": settings.GOOGLE_CLIENT_SECRET,
            },
            headers={"Content-Type": "application/x-www-form-urlencoded"},
        )
        if resp.status_code >= 400:
            raise KBProviderAuthError(
                f"Google token refresh failed: {resp.status_code}"
            )
        body = resp.json()
        self._access_token = body.get("access_token") or self._access_token
        new_refresh = body.get("refresh_token")
        if new_refresh:
            self._refresh_token = new_refresh
        if self._on_token_refresh is not None:
            try:
                await self._on_token_refresh(
                    self._access_token,
                    self._refresh_token,
                    body.get("expires_in"),
                )
            except Exception:
                logger.exception("on_token_refresh callback failed (gdrive)")


# ── Binary → text helpers ────────────────────────────────────────────


def _pdf_to_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        logger.warning("pypdf not installed; skipping PDF body")
        return ""
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _docx_to_text(data: bytes) -> str:
    try:
        import docx  # type: ignore
    except ImportError:
        logger.warning("python-docx not installed; skipping DOCX body")
        return ""
    document = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in document.paragraphs)


def _parse_rfc3339(raw: Optional[str]) -> Optional[datetime]:
    if not raw:
        return None
    try:
        return datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError:
        return None
