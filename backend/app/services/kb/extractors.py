"""Text extraction from uploaded KB files (PDF, DOCX, plain text).

Each extractor returns a unicode string. When extraction fails we raise
``ExtractionError`` so the caller can surface an actionable 400 to the user.
"""

from __future__ import annotations

import io
import logging
from typing import Callable, Dict

logger = logging.getLogger(__name__)


class ExtractionError(RuntimeError):
    """Raised when we fail to extract text from a file."""


def _extract_txt(data: bytes) -> str:
    # UTF-16 is deliberately omitted — on short byte streams it can silently
    # "succeed" on latin-1 input and emit garbage. UTF-8 then latin-1 covers
    # the realistic cases without false positives.
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return data.decode("latin-1")
        except UnicodeDecodeError as exc:
            raise ExtractionError("Unsupported text encoding") from exc


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover — installed in requirements
        raise ExtractionError("pypdf is not installed") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not parse PDF: {exc}") from exc

    if reader.is_encrypted:
        # Attempt an empty-password decrypt; many tools encrypt with no password.
        try:
            reader.decrypt("")
        except Exception:
            raise ExtractionError("PDF is password-protected")

    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            logger.warning("Skipping PDF page %d — extraction raised", i)
            pages.append("")

    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text:
        raise ExtractionError(
            "PDF contained no extractable text (likely a scanned image). "
            "Run it through OCR before uploading."
        )
    return text


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("python-docx is not installed") from exc

    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not parse DOCX: {exc}") from exc

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    text = "\n\n".join(parts)
    if not text.strip():
        raise ExtractionError("DOCX contained no extractable text")
    return text


_EXTENSION_MAP: Dict[str, Callable[[bytes], str]] = {
    ".txt": _extract_txt,
    ".md": _extract_txt,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def extract_text(filename: str, data: bytes) -> str:
    """Extract text from a KB upload, dispatching on filename extension.

    Falls back to UTF-8 text decode when the extension is unknown, so agents
    can upload .log, .json, or .rst snippets without us maintaining a
    dedicated extractor per format.
    """
    lower = (filename or "").lower()
    for suffix, extractor in _EXTENSION_MAP.items():
        if lower.endswith(suffix):
            return extractor(data)
    return _extract_txt(data)
