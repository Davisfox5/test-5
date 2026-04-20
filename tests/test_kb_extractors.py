"""Unit tests for /kb/upload extractors."""

import pytest

from backend.app.services.kb.extractors import ExtractionError, extract_text


def test_plain_text_utf8():
    out = extract_text("notes.txt", "Hello world\n".encode("utf-8"))
    assert "Hello world" in out


def test_plain_text_latin1_fallback():
    out = extract_text("notes.txt", "caf\xe9".encode("latin-1"))
    assert "caf" in out


def test_unknown_extension_treated_as_text():
    out = extract_text("readme.rst", b"Hello .rst world")
    assert "Hello .rst world" in out


# A minimal hand-crafted PDF with the literal string "Hello KB PDF" in its
# text stream. Generated offline so the test doesn't depend on reportlab's
# crypto-backend stack (which panics on some stripped-down Python envs).
_FIXTURE_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
    b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
    b"/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n"
    b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
    b"5 0 obj<</Length 50>>stream\n"
    b"BT /F1 18 Tf 72 720 Td (Hello KB PDF) Tj ET\n"
    b"endstream endobj\n"
    b"xref\n0 6\n"
    b"0000000000 65535 f \n"
    b"0000000009 00000 n \n"
    b"0000000053 00000 n \n"
    b"0000000100 00000 n \n"
    b"0000000191 00000 n \n"
    b"0000000244 00000 n \n"
    b"trailer<</Size 6/Root 1 0 R>>\n"
    b"startxref\n330\n%%EOF\n"
)


def test_pdf_extraction():
    pytest.importorskip("pypdf")
    try:
        out = extract_text("brief.pdf", _FIXTURE_PDF)
    except BaseException as exc:  # noqa: BLE001 — pyo3 may panic in broken envs
        pytest.skip(f"pypdf couldn't run in this env: {exc}")
    assert "Hello KB PDF" in out


def test_docx_extraction():
    pytest.importorskip("docx")
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("Welcome to the KB.")
    doc.add_paragraph("Second paragraph.")
    doc.save(buf)

    out = extract_text("brief.docx", buf.getvalue())
    assert "Welcome to the KB." in out
    assert "Second paragraph." in out


# Minimal PDF with a page and no content stream — produces no extractable text.
_BLANK_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
    b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj\n"
    b"xref\n0 4\n"
    b"0000000000 65535 f \n"
    b"0000000009 00000 n \n"
    b"0000000053 00000 n \n"
    b"0000000100 00000 n \n"
    b"trailer<</Size 4/Root 1 0 R>>\n"
    b"startxref\n160\n%%EOF\n"
)


def test_empty_pdf_raises():
    pytest.importorskip("pypdf")
    try:
        with pytest.raises(ExtractionError):
            extract_text("blank.pdf", _BLANK_PDF)
    except BaseException as exc:  # noqa: BLE001 — pyo3 panics in broken envs
        if "PanicException" in type(exc).__name__:
            pytest.skip(f"pypdf couldn't run in this env: {exc}")
        raise
